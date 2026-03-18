#!/usr/bin/env python3
"""
Polymarket full trade history export and PNL analysis.

Fetches complete trade + redeem history via Polymarket Data API (and optional
on-chain), maps usernames to proxy wallets, enriches with market category,
and writes per-trader CSVs plus a summary. Designed to match official
Polymarket/PolymarketAnalytics stats (PNL, positions, win rate).

Usage:
  pip install -r requirements.txt
  python polymarket_trader_history.py

  Optional: set POLYMARKET_URLS_FILE to a .txt path (one URL per line), or
  place polymarkettraders.docx in this folder to load URLs from the Word doc.
"""

from __future__ import annotations

import csv
import io
import re
import time
from pathlib import Path
from typing import Any, Callable

import pandas as pd
import requests

# -----------------------------------------------------------------------------
# Config
# -----------------------------------------------------------------------------
DATA_API_BASE = "https://data-api.polymarket.com"
GAMMA_API_BASE = "https://gamma-api.polymarket.com"
ACTIVITY_PAGE_SIZE = 500
TRADES_PAGE_SIZE = 10_000
RATE_LIMIT_SLEEP = 0.5  # seconds between API calls
OUTPUT_DIR = Path(__file__).resolve().parent / "output"
KNOWN_ADDRESSES: dict[str, str] = {
    "kch123": "0x6a72f61820b26b1fe4d956e17b6dc2a1ea3033ee",
    # Add more username -> proxy wallet (0x40 hex) as needed.
    # For usernames like @0xD9E0AACa471f48F91A26E8669A805f2 (short 0x) look up
    # on https://polymarketanalytics.com and add the full address here.
}

# Default list of profile URLs (deduplicated). Overridden by file if present.
DEFAULT_URLS = [
    "https://polymarket.com/@S-Works",
    "https://polymarket.com/@0xD9E0AACa471f48F91A26E8669A805f2",
    "https://polymarket.com/@Avarice31",
    "https://polymarket.com/@kch123",
    "https://polymarket.com/@BoomLaLa",
    "https://polymarket.com/@Bienville",
    "https://polymarket.com/@0x2c335066FE58fe9237c3d3Dc7b275C2a034a0563-1759935795465",
    "https://polymarket.com/@ckw",
    "https://polymarket.com/@ShortFlutterStock",
    "https://polymarket.com/@IBOV200K",
    "https://polymarket.com/@tcp2",
    "https://polymarket.com/@RandomPunter",
    "https://polymarket.com/@0xheavy888",
    "https://polymarket.com/@fkgggg2",
    "https://polymarket.com/@EIf",
    "https://polymarket.com/@LynxTitan",
    "https://polymarket.com/@bigmoneyloser00",
    "https://polymarket.com/@Capman",
    "https://polymarket.com/@geniusMC",
    "https://polymarket.com/@DLEK",
    "https://polymarket.com/@ShucksIt69",
    "https://polymarket.com/@0x20D6436849F930584892730C7F96eBB2Ac763856-1768642056357",
    "https://polymarket.com/@redskinrick",
    "https://polymarket.com/@JhonAlexanderHinestroza",
    "https://polymarket.com/@middleoftheocean",
    "https://polymarket.com/@Andromeda1",
    "https://polymarket.com/@CoryLahey",
    "https://polymarket.com/@TutiFromFactsOfLife",
    "https://polymarket.com/@TheArena",
    "https://polymarket.com/@xytest",
    "https://polymarket.com/@JuniorB",
    "https://polymarket.com/@UAEVALORANTFAN",
    "https://polymarket.com/@TheMangler",
    "https://polymarket.com/@iDropMyHotdog",
    "https://polymarket.com/@bloodmaster",
    "https://polymarket.com/@9sh8f",
    "https://polymarket.com/@0x53eCc53E7",
    "https://polymarket.com/@877s8d8g89I9f8d98fd99ww2",
    "https://polymarket.com/@Vetch",
    "https://polymarket.com/@TTdes",
]


def _sleep() -> None:
    time.sleep(RATE_LIMIT_SLEEP)


# -----------------------------------------------------------------------------
# URL parsing and username extraction
# -----------------------------------------------------------------------------
def parse_profile_url(url: str) -> str | None:
    """Extract username from Polymarket profile URL (part after @)."""
    url = (url or "").strip()
    m = re.match(r"https?://(?:www\.)?polymarket\.com/@(.+)", url, re.I)
    return m.group(1).strip() if m else None


def normalize_username(username: str) -> str:
    """Normalize for display and mapping (lowercase for 0x addresses)."""
    s = username.strip()
    if s.startswith("0x") and len(s) == 42 and all(c in "0123456789abcdefABCDEF" for c in s[2:]):
        return s.lower()
    if "-" in s and s.startswith("0x"):
        addr, _ = s.split("-", 1)
        if len(addr) == 42:
            return addr.lower()
    return s


def is_full_eth_address(s: str) -> bool:
    """True if s looks like 0x + 40 hex chars."""
    s = (s or "").strip()
    return bool(s.startswith("0x") and len(s) == 42 and re.match(r"0x[a-fA-F0-9]{40}", s))


def extract_address_from_username(username: str) -> str | None:
    """If username is 0x...-number, return the 0x40-char address part."""
    s = username.strip()
    if s.startswith("0x") and "-" in s:
        addr = s.split("-", 1)[0].strip()
        if len(addr) == 42 and re.match(r"0x[a-fA-F0-9]{40}", addr):
            return addr.lower()
    return None


def load_urls_from_file(path: Path) -> list[str]:
    """Load profile URLs from a text file (one URL per line)."""
    urls: list[str] = []
    text = path.read_text(encoding="utf-8", errors="replace")
    for line in text.splitlines():
        line = line.strip()
        if line and ("polymarket.com" in line and "@" in line):
            urls.append(line)
    return urls


def load_urls_from_docx(path: Path) -> list[str]:
    """Load profile URLs from a Word document (paragraphs and hyperlinks)."""
    try:
        from docx import Document
    except ImportError:
        return []
    urls: list[str] = []
    doc = Document(path)
    pattern = re.compile(r"https?://(?:www\.)?polymarket\.com/@[^\s\)\]\"]+")
    for p in doc.paragraphs:
        urls.extend(pattern.findall(p.text))
    for rel in getattr(doc, "part", None) and getattr(doc.part, "rels", []) or []:
        if getattr(rel, "target_ref", None) and "polymarket.com" in str(rel.target_ref):
            urls.append(str(rel.target_ref))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                urls.extend(pattern.findall(cell.text))
    return urls


def get_profile_urls() -> list[str]:
    """Collect profile URLs: from env file, docx in cwd, or default list. Dedupe."""
    seen: set[str] = set()
    out: list[str] = []
    base = Path(__file__).resolve().parent

    # Optional: explicit URL file
    import os
    url_file = os.environ.get("POLYMARKET_URLS_FILE")
    if url_file:
        p = Path(url_file)
        if p.is_file():
            for u in load_urls_from_file(p):
                u = u.strip()
                if u and u not in seen:
                    seen.add(u)
                    out.append(u)

    # Docx in project folder
    for name in ("polymarkettraders.docx", "polymarket_traders.docx"):
        p = base / name
        if p.is_file():
            for u in load_urls_from_docx(p):
                u = u.strip().split(")")[0].split("]")[0]  # trim trailing )
                if u and u not in seen:
                    seen.add(u)
                    out.append(u)
            break

    if not out:
        for u in DEFAULT_URLS:
            u = (u or "").strip()
            if u and u not in seen:
                seen.add(u)
                out.append(u)
    return out


# -----------------------------------------------------------------------------
# Username -> proxy wallet resolution
# -----------------------------------------------------------------------------
def resolve_username_to_address(username: str) -> str | None:
    """
    Resolve Polymarket username to Polygon proxy wallet (0x40 hex).
    - Full 0x40-char usernames: use as address.
    - 0x...-number: use address part before hyphen.
    - Else: Gamma public-search, then KNOWN_ADDRESSES.
    """
    raw = username.strip()
    if is_full_eth_address(raw):
        return raw.lower()
    addr = extract_address_from_username(raw)
    if addr:
        return addr
    # Known mapping
    key = raw.lower() if raw.startswith("0x") else raw
    if key in KNOWN_ADDRESSES:
        return KNOWN_ADDRESSES[key].lower()
    # Gamma API search (profile name match)
    try:
        r = requests.get(
            f"{GAMMA_API_BASE}/public-search",
            params={"q": raw, "search_profiles": "true", "limit_per_type": 20},
            timeout=15,
        )
        _sleep()
        r.raise_for_status()
        data = r.json()
        for profile in data.get("profiles") or []:
            if (profile.get("name") or "").strip() == raw or (profile.get("pseudonym") or "").strip() == raw:
                w = (profile.get("proxyWallet") or "").strip()
                if w and len(w) == 42:
                    return w.lower()
        # Fallback: first profile whose name starts with or contains the query
        for profile in data.get("profiles") or []:
            w = (profile.get("proxyWallet") or "").strip()
            if w and len(w) == 42 and raw.lower() in (profile.get("name") or "").lower():
                return w.lower()
    except Exception as e:
        print(f"  [resolve] Gamma search failed for '{username}': {e}")
    return None


# -----------------------------------------------------------------------------
# Data API: full activity and trades (paginated)
# -----------------------------------------------------------------------------
def fetch_all_activity(
    address: str,
    on_progress: Callable[[int, int], None] | None = None,
) -> list[dict[str, Any]]:
    """
    Paginate Data API activity until no more. Uses time-window pagination (end=timestamp)
    so we can get full history; the API caps offset at 10,000 so offset-only would
    miss data for large accounts.

    on_progress: optional callback(total_rows_so_far: int, page_number: int) after each batch.
    """
    address = address.lower()
    combined: list[dict[str, Any]] = []
    # Newest first (API default), then we request older batches with end = oldest_ts - 1
    end_ts: int | None = None
    page = 0
    while True:
        try:
            params: dict[str, Any] = {
                "user": address,
                "limit": ACTIVITY_PAGE_SIZE,
                "sortBy": "TIMESTAMP",
                "sortDirection": "DESC",
            }
            if end_ts is not None:
                params["end"] = end_ts
            r = requests.get(f"{DATA_API_BASE}/activity", params=params, timeout=30)
            _sleep()
            r.raise_for_status()
            batch = r.json()
        except requests.HTTPError as e:
            if e.response is not None and e.response.status_code == 400 and end_ts is None:
                # Fallback: some API versions may not like end on first request
                end_ts = 0
                continue
            print(f"  [activity] error at page {page}: {e}")
            break
        except Exception as e:
            print(f"  [activity] error at page {page}: {e}")
            break
        if not batch:
            break
        combined.extend(batch)
        if on_progress is not None:
            on_progress(len(combined), page + 1)
        if len(batch) < ACTIVITY_PAGE_SIZE:
            break
        oldest = min(b.get("timestamp") or 0 for b in batch)
        end_ts = oldest - 1
        page += 1
    return combined


def fetch_all_positions(address: str) -> list[dict[str, Any]]:
    """
    Fetch all current (open) positions for a user from Data API.
    Each position has cashPnl (unrealized), currentValue, initialValue, conditionId, etc.
    Paginate with limit 500 until no more.
    """
    address = address.lower()
    combined: list[dict[str, Any]] = []
    offset = 0
    limit = 500
    while True:
        try:
            r = requests.get(
                f"{DATA_API_BASE}/positions",
                params={"user": address, "limit": limit, "offset": offset},
                timeout=30,
            )
            _sleep()
            r.raise_for_status()
            batch = r.json()
        except Exception as e:
            print(f"  [positions] error at offset {offset}: {e}")
            break
        if not batch:
            break
        combined.extend(batch)
        if len(batch) < limit:
            break
        offset += limit
    return combined


def fetch_all_closed_positions(address: str) -> list[dict[str, Any]]:
    """
    Fetch all closed positions from Data API. Each has realizedPnl (Polymarket's canonical
    realized PnL for that position). Sum(realizedPnl) = official realized PnL, which matches
    analytics sites. API max limit is 50 per request.
    """
    address = address.lower()
    combined: list[dict[str, Any]] = []
    offset = 0
    limit = 50
    while True:
        try:
            r = requests.get(
                f"{DATA_API_BASE}/closed-positions",
                params={"user": address, "limit": limit, "offset": offset},
                timeout=30,
            )
            _sleep()
            r.raise_for_status()
            batch = r.json()
        except Exception as e:
            print(f"  [closed-positions] error at offset {offset}: {e}")
            break
        if not batch:
            break
        combined.extend(batch)
        if len(batch) < limit:
            break
        offset += limit
        if offset >= 100_000:
            break
    return combined


def fetch_all_trades(address: str) -> list[dict[str, Any]]:
    """Paginate Data API trades (takerOnly=false) to capture maker+taker. Merged with activity later."""
    address = address.lower()
    combined: list[dict[str, Any]] = []
    offset = 0
    while True:
        try:
            r = requests.get(
                f"{DATA_API_BASE}/trades",
                params={"user": address, "limit": TRADES_PAGE_SIZE, "offset": offset, "takerOnly": "false"},
                timeout=30,
            )
            _sleep()
            r.raise_for_status()
            batch = r.json()
        except Exception as e:
            print(f"  [trades] error at offset {offset}: {e}")
            break
        if not batch:
            break
        combined.extend(batch)
        if len(batch) < TRADES_PAGE_SIZE:
            break
        offset += TRADES_PAGE_SIZE
    return combined


# -----------------------------------------------------------------------------
# Build unified history from activity (primary) and optional trades
# -----------------------------------------------------------------------------
def activity_row_to_record(a: dict[str, Any], category: str = "") -> dict[str, Any]:
    """Turn one activity item into a flat record for CSV with analysis-friendly columns."""
    kind = (a.get("type") or "TRADE").upper()
    ts = a.get("timestamp") or 0
    size = float(a.get("size") or 0)
    usdc = float(a.get("usdcSize") or 0)
    price = float(a.get("price") or 0)
    slug = (a.get("slug") or "").strip()
    event_slug = (a.get("eventSlug") or "").strip()
    cat = category or infer_category_from_slug(slug or event_slug)
    submarket = infer_submarket_from_slug(slug, event_slug)
    cost_usdc = None
    payout_usdc = None
    if kind == "TRADE":
        cost_usdc = usdc
    elif kind == "REDEEM":
        payout_usdc = usdc
    elif kind in ("REWARD", "YIELD"):
        payout_usdc = usdc
    # Analysis columns: entry price and bet size for every row (filled for TRADE)
    entry_price = price if kind == "TRADE" else None
    bet_size_usdc = cost_usdc if kind == "TRADE" else None
    bet_size_shares = size if kind == "TRADE" else None
    return {
        "timestamp": ts,
        "datetime_utc": pd.Timestamp.utcfromtimestamp(ts).isoformat() if ts else "",
        "market_id": (a.get("conditionId") or "").strip(),
        "market_title": (a.get("title") or "").strip(),
        "slug": slug,
        "event_slug": event_slug,
        "category": cat,
        "submarket": submarket,
        "type": kind,
        "side": (a.get("side") or "").upper(),
        "outcome": (a.get("outcome") or "").strip(),
        "outcome_index": a.get("outcomeIndex") if isinstance(a.get("outcomeIndex"), (int, float)) else None,
        "size_shares": size,
        "size": size,
        "cost_usdc": cost_usdc,
        "payout_usdc": payout_usdc,
        "entry_price": entry_price,
        "bet_size_usdc": bet_size_usdc,
        "bet_size_shares": bet_size_shares,
        "price": price,
        "transaction_hash": (a.get("transactionHash") or "").strip(),
        "asset": (a.get("asset") or "").strip(),
        # Keep "title" as alias for market_title for backward compatibility
        "title": (a.get("title") or "").strip(),
    }


def infer_category_from_slug(slug: str) -> str:
    """Infer high-level category from slug (e.g. NHL, NBA, Politics)."""
    s = (slug or "").lower()
    if s.startswith("nhl-") or "nhl" in s[:10]:
        return "NHL"
    if s.startswith("nba-") or "nba" in s[:10]:
        return "NBA"
    if s.startswith("nfl-") or "nfl" in s[:10]:
        return "NFL"
    if s.startswith("mlb-") or "mlb" in s[:10]:
        return "MLB"
    if "politics" in s or "election" in s or "trump" in s or "biden" in s or "republican" in s or "presidential" in s:
        return "Politics"
    if "crypto" in s or "btc" in s or "bitcoin" in s or "eth" in s:
        return "Crypto"
    if "epl-" in s or "premier league" in s or "epl " in s:
        return "Soccer"
    if "ucl-" in s or "champions league" in s or "lal-" in s or "la liga" in s or "elc-" in s or "ere-" in s or "bun-" in s or "lol " in s:
        return "Soccer"
    if "sport" in s or "soccer" in s or "ufc" in s or "mma" in s:
        return "Sports"
    if "super bowl" in s or "superbowl" in s:
        return "NFL"
    if "cbb-" in s:
        return "NCAAB"
    return "Other"


def infer_submarket_from_slug(slug: str, event_slug: str = "") -> str:
    """Infer submarket for analysis: league, event type, or market theme."""
    s = ((slug or "") + " " + (event_slug or "")).lower()
    if not s.strip():
        return "Other"
    # Leagues / competitions
    if "epl-" in s or "epl " in s or "premier league" in s:
        return "EPL"
    if "ucl-" in s or "champions league" in s:
        return "UCL"
    if "lal-" in s or "la liga" in s:
        return "La Liga"
    if "elc-" in s:
        return "EFL Championship"
    if "ere-" in s:
        return "Eredivisie"
    if "bun-" in s:
        return "Bundesliga"
    if "nfl-" in s or "nfl " in s:
        return "NFL"
    if "super bowl" in s or "superbowl" in s:
        return "Super Bowl"
    if "nba-" in s or "nba " in s:
        return "NBA"
    if "nhl-" in s or "nhl " in s:
        return "NHL"
    if "mlb-" in s or "mlb " in s:
        return "MLB"
    if "cbb-" in s:
        return "NCAAB"
    if "lol " in s or "lol:" in s:
        return "LoL"
    if "presidential" in s or "election" in s:
        return "Elections"
    if "republican" in s or "nominee" in s:
        return "Politics"
    if "btc" in s or "bitcoin" in s:
        return "Crypto"
    return infer_category_from_slug(slug) or "Other"


# Optional: fetch category from Gamma by condition_id (cached to limit rate)
_category_cache: dict[str, str] = {}


def fetch_category_for_condition(condition_id: str) -> str:
    """Return category for condition_id from Gamma API (cached)."""
    if not condition_id or len(condition_id) != 66:
        return "Other"
    if condition_id in _category_cache:
        return _category_cache[condition_id]
    try:
        r = requests.get(
            f"{GAMMA_API_BASE}/markets",
            params={"condition_id": condition_id, "limit": 1},
            timeout=10,
        )
        _sleep()
        if r.ok:
            data = r.json()
            if isinstance(data, list) and data:
                m = data[0]
                cat = (m.get("category") or m.get("events", [{}])[0].get("category") or "").strip()
                if cat:
                    _category_cache[condition_id] = cat
                    return cat
    except Exception:
        pass
    _category_cache[condition_id] = "Other"
    return "Other"


def build_history_df(activity: list[dict[str, Any]], enrich_category: bool = False) -> pd.DataFrame:
    """Build one DataFrame from activity list; sort by timestamp. Optionally enrich category via Gamma."""
    if not activity:
        return pd.DataFrame()
    records = [activity_row_to_record(a) for a in activity]
    df = pd.DataFrame(records)
    df = df.sort_values("timestamp", ascending=True).reset_index(drop=True)
    if enrich_category and "market_id" in df.columns:
        for cid in df["market_id"].dropna().unique():
            if (cid and str(cid).startswith("0x") and
                    (df["market_id"] == cid).any() and (df.loc[df["market_id"] == cid, "category"] == "Other").any()):
                cat = fetch_category_for_condition(str(cid))
                df.loc[df["market_id"] == cid, "category"] = cat
    return df


# -----------------------------------------------------------------------------
# Metrics (PNL, ROI, win rate, positions, volume)
# -----------------------------------------------------------------------------
def analyze_history(df: pd.DataFrame) -> dict[str, Any]:
    """
    Compute PNL, ROI, win rate, total positions, volume to match official stats.
    - Positions: we count unique conditionIds that have a REDEEM (resolved positions).
    - Win rate: % of those positions where total redeem payout > total cost for that market.
    """
    if df.empty:
        return {
            "total_trades": 0,
            "total_redemptions": 0,
            "total_positions": 0,
            "volume_usdc": 0.0,
            "total_cost_usdc": 0.0,
            "total_sell_proceeds_usdc": 0.0,
            "total_redemption_usdc": 0.0,
            "other_income_usdc": 0.0,
            "other_cost_usdc": 0.0,
            "pnl_usdc": 0.0,
            "roi_pct": 0.0,
            "win_rate_pct": 0.0,
            "winning_positions": 0,
            "losing_positions": 0,
        }

    trades = df[df["type"] == "TRADE"]
    redeems = df[df["type"] == "REDEEM"]

    total_cost = float(trades.loc[trades["side"] == "BUY", "cost_usdc"].fillna(0).sum())
    total_sell = float(trades.loc[trades["side"] == "SELL", "cost_usdc"].fillna(0).sum())
    total_redemption = float(redeems["payout_usdc"].fillna(0).sum())

    # Only REWARD and YIELD as extra income (bonus USDC). MERGE/SPLIT excluded - position conversion, not PnL.
    other_income = float(df[df["type"].isin(["REWARD", "YIELD"])]["payout_usdc"].fillna(0).sum())
    other_cost = 0.0

    volume = float(trades["cost_usdc"].abs().sum())
    total_positions = redeems["market_id"].nunique()
    if total_positions == 0:
        total_positions = len(redeems)

    # Per-market P&L: cost (buys - sells) vs redemption
    market_cost: dict[str, float] = {}
    market_payout: dict[str, float] = {}
    for _, r in trades.iterrows():
        cid = r["market_id"]
        if cid not in market_cost:
            market_cost[cid] = 0.0
            market_payout[cid] = 0.0
        if r["side"] == "BUY":
            market_cost[cid] += float(r["cost_usdc"] or 0)
        else:
            market_cost[cid] -= float(r["cost_usdc"] or 0)
    for _, r in redeems.iterrows():
        cid = r["market_id"]
        if cid not in market_payout:
            market_payout[cid] = 0.0
        market_payout[cid] += float(r["payout_usdc"] or 0)

    winning = sum(1 for cid in market_payout if (market_payout.get(cid, 0) - market_cost.get(cid, 0)) > 0)
    losing = sum(1 for cid in market_payout if (market_payout.get(cid, 0) - market_cost.get(cid, 0)) <= 0)
    resolved_count = winning + losing
    win_rate = (100.0 * winning / resolved_count) if resolved_count else 0.0

    # Realized PNL: redemptions + sells - buys + REWARD + YIELD (matches Polymarket; MERGE/SPLIT excluded)
    pnl = total_redemption + total_sell - total_cost + other_income - other_cost
    total_invested = total_cost - total_sell
    roi = (100.0 * pnl / total_invested) if total_invested and total_invested > 0 else 0.0

    return {
        "total_trades": len(trades),
        "total_redemptions": len(redeems),
        "total_positions": total_positions,
        "volume_usdc": volume,
        "total_cost_usdc": total_cost,
        "total_sell_proceeds_usdc": total_sell,
        "total_redemption_usdc": total_redemption,
        "other_income_usdc": other_income,
        "other_cost_usdc": other_cost,
        "pnl_usdc": pnl,
        "realized_pnl_usdc": pnl,  # alias for clarity when we add unrealized
        "roi_pct": roi,
        "win_rate_pct": win_rate,
        "winning_positions": winning,
        "losing_positions": losing,
    }


def add_open_positions_metrics(
    metrics: dict[str, Any],
    positions: list[dict[str, Any]],
    closed_positions: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """
    Enrich metrics with open positions (unrealized PnL) and optionally canonical
    realized PnL from GET /closed-positions. When closed_positions is provided,
    realized_pnl_usdc = sum(realizedPnl) from API so total PnL matches Polymarket/analytics.
    """
    out = dict(metrics)
    # Canonical realized from Data API closed-positions (matches Polymarket/analytics when available)
    if closed_positions is not None:
        api_realized = sum(float(c.get("realizedPnl") or 0) for c in closed_positions)
        out["realized_pnl_usdc"] = api_realized
        out["pnl_usdc"] = api_realized
        out["_realized_source"] = "closed_positions_api"  # used by one-off script for console label only
    realized = out.get("pnl_usdc") or 0.0

    if not positions:
        out["open_positions_count"] = 0
        out["unrealized_pnl_usdc"] = 0.0
        out["total_pnl_usdc"] = realized
        out["total_predictions"] = out.get("total_positions") or 0
        return out
    unrealized = sum(float(p.get("cashPnl") or 0) for p in positions)
    open_by_condition = len(set((p.get("conditionId") or "").strip() for p in positions if (p.get("conditionId") or "").strip()))
    out["open_positions_count"] = open_by_condition
    out["unrealized_pnl_usdc"] = unrealized
    out["total_pnl_usdc"] = realized + unrealized
    out["total_predictions"] = (out.get("total_positions") or 0) + open_by_condition
    return out


# -----------------------------------------------------------------------------
# Positions summary (one row per closed position for strategy analysis)
# -----------------------------------------------------------------------------
def build_positions_summary_df(closed_positions: list[dict[str, Any]]) -> pd.DataFrame:
    """
    Build a DataFrame with one row per closed position: market, outcome, cost, realized PnL, win/loss.
    Use for analyzing which markets/submarkets the trader wins in and bet sizes.
    """
    if not closed_positions:
        return pd.DataFrame()
    rows = []
    for c in closed_positions:
        cid = (c.get("conditionId") or "").strip()
        slug = (c.get("slug") or "").strip()
        event_slug = (c.get("eventSlug") or "").strip()
        cat = infer_category_from_slug(slug or event_slug)
        submarket = infer_submarket_from_slug(slug, event_slug)
        realized = float(c.get("realizedPnl") or 0)
        total_bought = float(c.get("totalBought") or 0)
        avg_price = float(c.get("avgPrice") or 0)
        cur_price = float(c.get("curPrice") or 0)
        rows.append({
            "market_id": cid,
            "market_title": (c.get("title") or "").strip(),
            "slug": slug,
            "event_slug": event_slug,
            "category": cat,
            "submarket": submarket,
            "outcome": (c.get("outcome") or "").strip(),
            "outcome_index": c.get("outcomeIndex"),
            "total_bought_usdc": total_bought,
            "avg_entry_price": avg_price,
            "resolved_price": cur_price,
            "realized_pnl_usdc": realized,
            "win": realized > 0,
        })
    return pd.DataFrame(rows)


def _position_size_bucket(total_bought: float) -> str:
    """Bucket for ROI-by-size breakdown."""
    if total_bought < 1_000:
        return "<$1k"
    if total_bought < 10_000:
        return "$1k-$10k"
    if total_bought < 50_000:
        return "$10k-$50k"
    if total_bought < 100_000:
        return "$50k-$100k"
    if total_bought < 500_000:
        return "$100k-$500k"
    return "$500k+"


def build_breakdowns(positions_summary_df: pd.DataFrame) -> dict[str, pd.DataFrame]:
    """
    Build breakdowns of realized PnL by category, submarket, and position-size bucket.
    Returns dict with keys: by_category, by_submarket, by_position_size.
    Each DataFrame has: segment name, positions, total_bought_usdc, realized_pnl_usdc, roi_pct, wins, losses, win_rate_pct.
    """
    if positions_summary_df.empty or "realized_pnl_usdc" not in positions_summary_df.columns:
        return {"by_category": pd.DataFrame(), "by_submarket": pd.DataFrame(), "by_position_size": pd.DataFrame()}
    df = positions_summary_df.copy()
    df["position_size_bucket"] = df["total_bought_usdc"].fillna(0).apply(_position_size_bucket)

    def agg(g: pd.DataFrame) -> pd.DataFrame:
        total_bought = g["total_bought_usdc"].sum()
        realized = g["realized_pnl_usdc"].sum()
        wins = g["win"].sum()
        losses = len(g) - wins
        wr = (100.0 * wins / len(g)) if len(g) else 0.0
        roi = (100.0 * realized / total_bought) if total_bought and total_bought > 0 else 0.0
        return pd.DataFrame([{
            "positions": len(g),
            "total_bought_usdc": total_bought,
            "realized_pnl_usdc": realized,
            "roi_pct": round(roi, 2),
            "wins": int(wins),
            "losses": int(losses),
            "win_rate_pct": round(wr, 1),
        }])

    by_cat = df.groupby("category", dropna=False).apply(agg, include_groups=False).reset_index(level=1, drop=True).reset_index()
    by_cat = by_cat.rename(columns={"category": "segment"})
    by_sub = df.groupby("submarket", dropna=False).apply(agg, include_groups=False).reset_index(level=1, drop=True).reset_index()
    by_sub = by_sub.rename(columns={"submarket": "segment"})
    # Order position-size buckets logically
    bucket_order = ["<$1k", "$1k-$10k", "$10k-$50k", "$50k-$100k", "$100k-$500k", "$500k+"]
    by_size = df.groupby("position_size_bucket", sort=False).apply(agg, include_groups=False).reset_index(level=1, drop=True).reset_index()
    by_size = by_size.rename(columns={"position_size_bucket": "segment"})
    by_size["segment"] = pd.Categorical(by_size["segment"], categories=bucket_order, ordered=True)
    by_size = by_size.sort_values("segment").reset_index(drop=True)
    by_size["segment"] = by_size["segment"].astype(str)
    return {"by_category": by_cat, "by_submarket": by_sub, "by_position_size": by_size}


# -----------------------------------------------------------------------------
# CSV export
# -----------------------------------------------------------------------------
def save_to_csv(
    username: str,
    df: pd.DataFrame,
    metrics: dict[str, Any],
    positions_summary_df: pd.DataFrame | None = None,
) -> Path:
    """Write {username}_trades.csv (every trade/activity row) and optional {username}_positions_summary.csv. Returns path to trades CSV."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]", "_", username)[:80]
    path = OUTPUT_DIR / f"{safe_name}_trades.csv"
    df.to_csv(path, index=False, quoting=csv.QUOTE_MINIMAL)
    if positions_summary_df is not None and not positions_summary_df.empty:
        summary_csv_path = OUTPUT_DIR / f"{safe_name}_positions_summary.csv"
        positions_summary_df.to_csv(summary_csv_path, index=False, quoting=csv.QUOTE_MINIMAL)
    summary_path = OUTPUT_DIR / f"{safe_name}_summary.txt"
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(f"Trader: {username}\n\n")
        # PNL tie-out (verify against Polymarket / Polymarket Analytics)
        realized = metrics.get("realized_pnl_usdc", metrics.get("pnl_usdc", 0))
        unrealized = metrics.get("unrealized_pnl_usdc", 0)
        total_pnl = metrics.get("total_pnl_usdc", realized + unrealized)
        f.write("PNL TIE-OUT (check against Polymarket / Polymarket Analytics)\n")
        f.write("-" * 50 + "\n")
        f.write(f"Realized PNL (closed positions, from API):  {realized:,.2f}\n")
        f.write(f"Unrealized PNL (open positions):           {unrealized:,.2f}\n")
        f.write(f"Total PNL:                                 {total_pnl:,.2f}\n\n")
        f.write("KEY METRICS\n")
        f.write("-" * 50 + "\n")
        f.write(f"ROI (%): {metrics['roi_pct']:.2f}\n")
        f.write(f"Win rate (%): {metrics['win_rate_pct']:.1f}\n")
        f.write(f"Winning positions: {metrics['winning_positions']}\n")
        f.write(f"Losing positions: {metrics['losing_positions']}\n")
        if "avg_trade_size_usdc" in metrics:
            f.write(f"Average trade size (USDC, per BUY fill): {metrics['avg_trade_size_usdc']:,.2f}\n")
        if "median_trade_size_usdc" in metrics:
            f.write(f"Median trade size (USDC): {metrics['median_trade_size_usdc']:,.2f}\n")
        if "avg_position_size_usdc" in metrics:
            f.write(f"Average position size (USDC, per closed position): {metrics['avg_position_size_usdc']:,.2f}\n")
        f.write("\nCOUNTS & VOLUME\n")
        f.write("-" * 50 + "\n")
        f.write(f"Total trades: {metrics['total_trades']}\n")
        f.write(f"Total redemptions: {metrics['total_redemptions']}\n")
        f.write(f"Closed positions (resolved): {metrics['total_positions']}\n")
        if "open_positions_count" in metrics:
            f.write(f"Open positions: {metrics['open_positions_count']}\n")
            f.write(f"Total predictions (closed + open): {metrics.get('total_predictions', '')}\n")
        f.write(f"Volume (USDC): {metrics['volume_usdc']:,.2f}\n")
        f.write(f"Total cost (buys): {metrics['total_cost_usdc']:,.2f}\n")
        f.write(f"Sell proceeds: {metrics['total_sell_proceeds_usdc']:,.2f}\n")
        f.write(f"Redemption payouts: {metrics['total_redemption_usdc']:,.2f}\n")
        if metrics.get("other_income_usdc"):
            f.write(f"Other income (REWARD/YIELD): {metrics.get('other_income_usdc', 0):,.2f}\n")
    return path


def write_summary_csv(summaries: list[dict[str, Any]]) -> Path:
    """Write summary CSV of all traders (ranked by PNL)."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / "all_traders_summary.csv"
    df = pd.DataFrame(summaries)
    if not df.empty:
        df = df.sort_values("pnl_usdc", ascending=False).reset_index(drop=True)
    df.to_csv(path, index=False)
    return path


# -----------------------------------------------------------------------------
# Optional validation (e.g. known trader stats)
# -----------------------------------------------------------------------------
def validate_known_trader(username: str, metrics: dict[str, Any]) -> None:
    """Assert approximate match for known traders (e.g. kch123)."""
    if username != "kch123":
        return
    # Official: PNL $10,620,160, 3,190 positions, 54.1% win rate (from prompt)
    pnl = metrics.get("pnl_usdc") or 0
    pos = metrics.get("total_positions") or 0
    wr = metrics.get("win_rate_pct") or 0
    if abs(pnl - 10_620_160) > 0.02 * 10_620_160:
        print(f"  [validation] @kch123 PNL mismatch: got {pnl:,.0f}, expected ~10,620,160")
    if pos and abs(pos - 3190) > 0.05 * 3190:
        print(f"  [validation] @kch123 positions mismatch: got {pos}, expected ~3190")
    if wr and abs(wr - 54.1) > 5:
        print(f"  [validation] @kch123 win rate mismatch: got {wr:.1f}%, expected ~54.1%")


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
def get_history(address: str) -> list[dict[str, Any]]:
    """Fetch full history for one wallet (activity only; trades endpoint can be added)."""
    activity = fetch_all_activity(address)
    return activity


def process_trader(username: str, address: str, enrich_category: bool = False) -> dict[str, Any] | None:
    """Fetch history, build DF, compute metrics, save CSV. Returns summary dict or None on skip."""
    import os
    enrich = enrich_category or (os.environ.get("ENRICH_CATEGORY_FROM_GAMMA", "").lower() in ("1", "true", "yes"))
    print(f"  Fetching history for @{username} ({address[:10]}...)")
    activity = get_history(address)
    if not activity:
        print(f"  No activity for @{username}")
        return {"username": username, "address": address, "pnl_usdc": 0, "total_positions": 0, "win_rate_pct": 0, "volume_usdc": 0, "roi_pct": 0, "total_trades": 0, "total_redemptions": 0, "winning_positions": 0, "losing_positions": 0}
    df = build_history_df(activity, enrich_category=enrich)
    metrics = analyze_history(df)
    metrics["username"] = username
    metrics["address"] = address
    save_to_csv(username, df, metrics)
    validate_known_trader(username, metrics)
    print(f"    -> {len(df)} rows, PNL ${metrics['pnl_usdc']:,.0f}, positions {metrics['total_positions']}, win rate {metrics['win_rate_pct']:.1f}%")
    return metrics


def main() -> None:
    urls = get_profile_urls()
    usernames_raw: list[str] = []
    for u in urls:
        uname = parse_profile_url(u)
        if uname and uname not in usernames_raw:
            usernames_raw.append(uname)

    print(f"Loaded {len(usernames_raw)} unique usernames")

    resolved: list[tuple[str, str]] = []
    for username in usernames_raw:
        addr = resolve_username_to_address(username)
        if addr:
            resolved.append((username, addr))
        else:
            print(f"Skipping (no address): @{username} — add to KNOWN_ADDRESSES or find on polymarketanalytics.com")

    summaries: list[dict[str, Any]] = []
    for username, address in resolved:
        try:
            s = process_trader(username, address)
            if s:
                summaries.append(s)
        except Exception as e:
            print(f"  Error processing @{username}: {e}")
            summaries.append({"username": username, "address": address, "pnl_usdc": 0, "total_positions": 0, "win_rate_pct": 0, "volume_usdc": 0, "roi_pct": 0, "total_trades": 0, "total_redemptions": 0, "winning_positions": 0, "losing_positions": 0})

    if summaries:
        write_summary_csv(summaries)
        print("\n--- Rankings by PNL ---")
        for i, s in enumerate(sorted(summaries, key=lambda x: x.get("pnl_usdc", 0), reverse=True)[:20], 1):
            print(f"  {i}. @{s.get('username','?')}  PNL ${s.get('pnl_usdc',0):,.0f}  ROI {s.get('roi_pct',0):.1f}%  positions {s.get('total_positions',0)}  win rate {s.get('win_rate_pct',0):.1f}%")
        print(f"\nOutput dir: {OUTPUT_DIR}")
    else:
        print("No traders processed.")


if __name__ == "__main__":
    main()
